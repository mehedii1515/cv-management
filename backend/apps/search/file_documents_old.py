from django_elasticsearch_dsl import Document, Index, fields
from django.conf import settings
import os
import hashlib
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

# Create the file index
file_index = Index('file_index')
file_index.settings(
    number_of_shards=1,
    number_of_replicas=0,
    analysis={
        'analyzer': {
            'file_analyzer': {
                'type': 'standard',
                'stopwords': '_english_'
            }
        }
    }
)

# Pure File Document - Independent of Resume Database
@file_index.doc_type
class FileDocument(Document):
    """
    Pure file-based document for DTSearch-like file indexing
    This indexes actual files directly, independent of database records
    """
    
    # File Identity
    file_path = fields.KeywordField()           # Full path to file
    filename = fields.TextField()               # Filename with analysis
    file_extension = fields.KeywordField()      # File extension
    file_size = fields.IntegerField()          # File size in bytes
    file_hash = fields.KeywordField()          # File content hash
    
    # File Timestamps
    created_date = fields.DateField()          # File creation date
    modified_date = fields.DateField()         # File modification date  
    indexed_date = fields.DateField()          # When indexed by system
    
    # File Content - The main searchable content
    content = fields.TextField(
        analyzer='standard',
        search_analyzer='standard'
    )
    
    # File Metadata
    file_type_category = fields.KeywordField() # resume, document, etc.
    language = fields.KeywordField()           # Detected language
    page_count = fields.IntegerField()         # Number of pages (for PDFs)
    word_count = fields.IntegerField()         # Estimated word count
    
    # Directory Structure
    directory_path = fields.KeywordField()     # Directory containing file
    relative_path = fields.KeywordField()      # Relative path from base

    @classmethod
    def create_from_file(cls, file_path: str, base_directory: str = None):
        """
        Create a FileDocument from a physical file
        This is the main method to index files
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return None
                
            # Get file stats
            file_stats = os.stat(file_path)
            filename = os.path.basename(file_path)
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Extract text content
            content = cls.extract_text_content(file_path)
            
            # Generate file hash
            file_hash = cls.generate_file_hash(file_path)
            
            # Calculate relative path
            relative_path = file_path
            directory_path = os.path.dirname(file_path)
            if base_directory and file_path.startswith(base_directory):
                relative_path = os.path.relpath(file_path, base_directory)
                
            # Create document
            doc = cls(
                meta={'id': file_hash},  # Use file hash as unique ID
                file_path=file_path,
                filename=filename,
                file_extension=file_extension,
                file_size=file_stats.st_size,
                file_hash=file_hash,
                created_date=datetime.fromtimestamp(file_stats.st_ctime),
                modified_date=datetime.fromtimestamp(file_stats.st_mtime),
                indexed_date=datetime.now(),
                content=content,
                file_type_category=cls.determine_file_category(file_path, content),
                language=cls.detect_language(content),
                page_count=cls.get_page_count(file_path, file_extension),
                word_count=len(content.split()) if content else 0,
                directory_path=directory_path,
                relative_path=relative_path
            )
            
            return doc
            
        except Exception as e:
            logger.error(f"Error creating FileDocument from {file_path}: {e}")
            return None

    @staticmethod
    def extract_text_content(file_path: str) -> str:
        """Extract text content from various file types"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return FileDocument.extract_pdf_content(file_path)
            elif file_extension == '.docx':
                return FileDocument.extract_docx_content(file_path)
            elif file_extension == '.doc':
                return FileDocument.extract_doc_content(file_path)
            elif file_extension in ['.txt', '.text']:
                return FileDocument.extract_txt_content(file_path)
            elif file_extension in ['.rtf']:
                return FileDocument.extract_rtf_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return ""

    @staticmethod
    def extract_pdf_content(file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF content from {file_path}: {e}")
            return ""

    @staticmethod
    def extract_docx_content(file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return "\n".join(text)
        except Exception as e:
            # Fallback to unstructured
            try:
                from unstructured.partition.docx import partition_docx
                elements = partition_docx(filename=file_path)
                return "\n".join([str(element) for element in elements])
            except Exception as e2:
                logger.error(f"Error extracting DOCX content from {file_path}: {e}, {e2}")
                return ""

    @staticmethod
    def extract_doc_content(file_path: str) -> str:
        """Extract text from DOC files"""
        try:
            from unstructured.partition.doc import partition_doc
            elements = partition_doc(filename=file_path)
            return "\n".join([str(element) for element in elements])
        except Exception as e:
            logger.error(f"Error extracting DOC content from {file_path}: {e}")
            return ""

    @staticmethod
    def extract_txt_content(file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {e}")
                return ""

    @staticmethod
    def extract_rtf_content(file_path: str) -> str:
        """Extract text from RTF files"""
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            logger.error(f"Error extracting RTF content from {file_path}: {e}")
            return ""

    @staticmethod
    def generate_file_hash(file_path: str) -> str:
        """Generate MD5 hash of file content"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Error generating hash for {file_path}: {e}")
            return ""

    @staticmethod
    def determine_file_category(file_path: str, content: str) -> str:
        """Determine the category of file based on path and content"""
        filename_lower = os.path.basename(file_path).lower()
        content_lower = content.lower() if content else ""
        
        # Resume indicators
        resume_indicators = ['resume', 'cv', 'curriculum', 'vitae']
        if any(indicator in filename_lower for indicator in resume_indicators):
            return 'resume'
            
        # Check content for resume-like patterns
        resume_content_patterns = [
            'experience', 'education', 'skills', 'employment',
            'objective', 'summary', 'qualifications', 'achievements'
        ]
        if sum(1 for pattern in resume_content_patterns if pattern in content_lower) >= 3:
            return 'resume'
            
        # Cover letter
        if 'cover' in filename_lower and 'letter' in filename_lower:
            return 'cover_letter'
            
        # Portfolio
        if 'portfolio' in filename_lower:
            return 'portfolio'
            
        return 'document'

    @staticmethod
    def detect_language(content: str) -> str:
        """Basic language detection"""
        if not content:
            return 'unknown'
            
        # Simple heuristic - could be enhanced with proper language detection
        english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for']
        content_lower = content.lower()
        english_count = sum(1 for word in english_words if word in content_lower)
        
        return 'english' if english_count >= 3 else 'other'

    @staticmethod
    def get_page_count(file_path: str, file_extension: str) -> int:
        """Get page count for supported file types"""
        try:
            if file_extension == '.pdf':
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                return len(reader.pages)
            elif file_extension == '.docx':
                from docx import Document
                doc = Document(file_path)
                # Rough estimate - DOCX doesn't have explicit page count
                return max(1, len(doc.paragraphs) // 20)
        except:
            pass
        return 1
